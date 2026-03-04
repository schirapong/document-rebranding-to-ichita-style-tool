#!/usr/bin/env python3
"""
Convert Markdown documents to Ichita-branded .docx files.

Applies Ichita Brand Identity styling:
- Colors: Blue #2978FF (accent), Blue Grey 03 #263338 (headings/body),
          Blue Light #82B0FF, Blue Grey 01 #CFD9DB (backgrounds)
- Font: Aeonik (with Arial Narrow / Calibri fallback if not installed)
- Logo: ICHITA wordmark in header
- Footer: ichitaglobal.com | Separation Technologies
- Professional, technical, innovative tone

Handles: headings (#-####), tables, code blocks, bold/italic, bullets,
numbered lists, blockquotes, horizontal rules, links, and regular paragraphs.
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Ichita Brand Colours ─────────────────────────────────────────────────────

ICHITA_BLUE       = RGBColor(0x29, 0x78, 0xFF)  # #2978FF — primary accent
ICHITA_BLUE_LIGHT = RGBColor(0x82, 0xB0, 0xFF)  # #82B0FF — light accent
ICHITA_BLUE_GREY1 = RGBColor(0xCF, 0xD9, 0xDB)  # #CFD9DB — light background
ICHITA_BLUE_GREY2 = RGBColor(0x78, 0x8F, 0x9C)  # #788F9C — medium grey
ICHITA_BLUE_GREY3 = RGBColor(0x26, 0x33, 0x38)  # #263338 — dark text, headings
ICHITA_BLUE_BLACK = RGBColor(0x17, 0x1C, 0x21)  # #171C21 — deepest dark
WHITE             = RGBColor(0xFF, 0xFF, 0xFF)

# Hex strings for cell shading (no # prefix)
HEX_BLUE         = "2978FF"
HEX_BLUE_LIGHT   = "82B0FF"
HEX_BLUE_GREY1   = "CFD9DB"
HEX_BLUE_GREY3   = "263338"
HEX_TABLE_HDR    = "263338"   # Dark header row
HEX_TABLE_ALT    = "EFF2F3"   # Very light alternating row (subtle Blue Grey 01 tint)
HEX_CODE_BG      = "EFF2F3"   # Code block background
HEX_QUOTE_BG     = "EBF0F7"   # Blockquote background


# ── Font Configuration ────────────────────────────────────────────────────────
# Priority: Aeonik (brand) → Avenir Next (closest geometric match) → Calibri (fallback)

BRAND_FONT = "Avenir Next"  # Closest system match to Aeonik
THAI_FONT  = "Bai Jamjuree" # Thai font with matched metrics to geometric sans-serif
THAI_SCALE = 0.9             # Thai 9pt / English 10pt — Bai Jamjuree one size down for visual balance
MONO_FONT  = "Courier New"

# Try to detect Aeonik availability
for font_dir in [os.path.expanduser("~/Library/Fonts"), "/Library/Fonts", "/System/Library/Fonts"]:
    if os.path.isdir(font_dir):
        for f in os.listdir(font_dir):
            if "aeonik" in f.lower():
                BRAND_FONT = "Aeonik"
                break

# ── Logo Path ─────────────────────────────────────────────────────────────────
# Logo-05: dark wordmark on white/transparent background — ideal for headers
# Resolve relative to script location so it works from any working directory
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(_SCRIPT_DIR, "assets", "Ichita_Logo-05.png")
if not os.path.exists(LOGO_PATH):
    # Fallback to absolute path on this machine
    LOGO_PATH = os.path.expanduser(
        "~/Documents/my-project/ichita brand ID/"
        "ICHITA BRAND BOOK AND COMPANY PROFILE/LogoV2 3/Digital/PNG/"
        "Wordmark/Ichita_Logo-05.png"
    )


# ── Helpers ───────────────────────────────────────────────────────────────────

def _is_thai(c):
    """Check if a character is Thai (U+0E00-U+0E7F)."""
    return '\u0e00' <= c <= '\u0e7f'


def _split_thai_latin(text):
    """Split text into segments of (text, is_thai) tuples.
    Groups consecutive Thai chars together, and consecutive non-Thai chars together."""
    if not text:
        return []
    segments = []
    current = text[0]
    current_thai = _is_thai(text[0])
    for c in text[1:]:
        c_thai = _is_thai(c)
        if c_thai == current_thai:
            current += c
        else:
            segments.append((current, current_thai))
            current = c
            current_thai = c_thai
    segments.append((current, current_thai))
    return segments


def _add_split_run(paragraph, text, font_name, base_size, color, bold=False,
                   italic=False, underline=False, is_link=False):
    """Add text as split Thai/Latin runs with different fonts and sizes.
    Thai segments get Bai Jamjuree at scaled size, Latin gets Aeonik at base size."""
    thai_size = Pt(round(base_size / 12700 * THAI_SCALE * 2) / 2)  # EMU → pt, scale, back to EMU
    for segment, is_thai in _split_thai_latin(text):
        run = paragraph.add_run(segment)
        run.font.name = THAI_FONT if is_thai else font_name
        run.font.size = thai_size if is_thai else base_size
        run.font.color.rgb = color
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if underline:
            run.font.underline = True
    return run if text else None

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color="788F9C"):
    """Apply borders to every cell in the table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_text(paragraph, text, base_font=None, base_size=Pt(10),
                       base_color=None, is_blockquote=False):
    """Parse inline markdown (bold, italic, bold+italic, links) and add runs.
    Splits Thai/Latin into separate runs with matched visual sizes."""
    if base_font is None:
        base_font = BRAND_FONT
    if base_color is None:
        base_color = ICHITA_BLUE_GREY3

    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'       # bold+italic
        r'|(\*\*(.+?)\*\*)'           # bold
        r'|(\*(.+?)\*)'               # italic
        r'|(\[([^\]]+)\]\(([^)]+)\))' # link
    )

    last_end = 0
    for match in pattern.finditer(text):
        before = text[last_end:match.start()]
        if before:
            _add_split_run(paragraph, before, base_font, base_size, base_color,
                           italic=is_blockquote)

        if match.group(2):  # ***bold+italic***
            _add_split_run(paragraph, match.group(2), base_font, base_size,
                           base_color, bold=True, italic=True)
        elif match.group(4):  # **bold**
            _add_split_run(paragraph, match.group(4), base_font, base_size,
                           base_color, bold=True, italic=is_blockquote)
        elif match.group(6):  # *italic*
            _add_split_run(paragraph, match.group(6), base_font, base_size,
                           base_color, italic=True)
        elif match.group(8):  # [link](url)
            _add_split_run(paragraph, match.group(9), base_font, base_size,
                           ICHITA_BLUE, underline=True)
        last_end = match.end()

    remaining = text[last_end:]
    if remaining:
        _add_split_run(paragraph, remaining, base_font, base_size, base_color,
                       italic=is_blockquote)


def add_cell_formatted_text(cell, text, is_header=False, font_name=None, font_size=Pt(9)):
    """Add formatted text to a table cell, preserving bold markdown."""
    if font_name is None:
        font_name = BRAND_FONT
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)

    if is_header:
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        _add_split_run(paragraph, clean, font_name, font_size, WHITE, bold=True)
    else:
        pattern = re.compile(r'\*\*(.+?)\*\*')
        last_end = 0
        for match in pattern.finditer(text):
            before = text[last_end:match.start()]
            if before:
                _add_split_run(paragraph, before, font_name, font_size, ICHITA_BLUE_GREY3)
            _add_split_run(paragraph, match.group(1), font_name, font_size,
                           ICHITA_BLUE_GREY3, bold=True)
            last_end = match.end()
        remaining = text[last_end:]
        if remaining:
            _add_split_run(paragraph, remaining, font_name, font_size, ICHITA_BLUE_GREY3)


def parse_table_line(line):
    """Parse a markdown table line into cell values."""
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]


def is_separator_line(line):
    """Check if a line is a markdown table separator (|---|---|)."""
    stripped = line.strip().strip('|')
    parts = stripped.split('|')
    return all(re.match(r'^[\s\-:]+$', p) for p in parts)


def add_code_block(doc, code_lines):
    """Add a code block with monospace font and branded background.
    Uses keep_together to prevent page breaks within the block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.keep_together = True

    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_CODE_BG}" w:val="clear"/>')
    pPr.append(shading)

    code_text = '\n'.join(code_lines)
    run = p.add_run(code_text)
    run.font.name = MONO_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = ICHITA_BLUE_GREY3


def add_table_from_rows(doc, header_cells, data_rows):
    """Create a branded table from parsed markdown table header and data rows."""
    num_cols = len(header_cells)

    normalized_data = []
    for row in data_rows:
        while len(row) < num_cols:
            row.append("")
        normalized_data.append(row[:num_cols])

    table = doc.add_table(rows=1 + len(normalized_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    set_table_borders(table, color="A0B0B8")

    # Prevent rows from splitting across pages and keep table together
    for row in table.rows:
        # cantSplit: prevents a single row from breaking across pages
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        # keep_together on each cell paragraph to keep table on one page
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = True
    # Remove keep_with_next from last row (nothing to keep with after table)
    last_row = table.rows[-1]
    for cell in last_row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = False

    # Header row — dark Ichita Blue Grey 03 background, white text
    hdr_row = table.rows[0]
    for j, cell_text in enumerate(header_cells):
        if j < num_cols:
            add_cell_formatted_text(hdr_row.cells[j], cell_text, is_header=True)
            set_cell_shading(hdr_row.cells[j], HEX_TABLE_HDR)

    # Data rows with subtle alternating shading
    for row_idx, row_data in enumerate(normalized_data):
        row = table.rows[row_idx + 1]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                add_cell_formatted_text(row.cells[j], cell_text)
                if row_idx % 2 == 1:
                    set_cell_shading(row.cells[j], HEX_TABLE_ALT)

    # Small spacing after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)


def add_header_footer(doc, logo_path):
    """Add ICHITA logo to header and branded footer to all sections."""
    for section in doc.sections:
        # ── Header ──
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing
        for p in header.paragraphs:
            p.clear()

        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(4)

        # Add logo image
        if os.path.exists(logo_path):
            run = hp.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
        else:
            # Fallback: text-based logo
            run = hp.add_run("ICHITA\u2122")
            run.font.name = BRAND_FONT
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = ICHITA_BLUE_GREY3

        # Add thin blue line under header
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="2978FF"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # ── Footer — disabled ──
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()


def add_title_page_band(doc):
    """Add a subtle Ichita Blue accent band (paragraph) after the title."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    # Thin blue bottom border as a decorative band
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="24" w:space="1" w:color="2978FF"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ── Main Conversion ──────────────────────────────────────────────────────────

def convert_md_to_docx(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ── Default document font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = BRAND_FONT
    font.size = Pt(10)
    font.color.rgb = ICHITA_BLUE_GREY3
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.space_before = Pt(2)
    # Set Thai (Complex Script) font + scaled size on default style
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), THAI_FONT)
    rFonts.set(qn('w:eastAsia'), THAI_FONT)
    scaled_hp = str(round(10 * 2 * THAI_SCALE))
    szCs_el = rPr.find(qn('w:szCs'))
    if szCs_el is not None:
        szCs_el.set(qn('w:val'), scaled_hp)
    else:
        szCs_el = OxmlElement('w:szCs')
        szCs_el.set(qn('w:val'), scaled_hp)
        rPr.append(szCs_el)

    # ── Heading styles (Ichita branding) ──
    heading_configs = [
        (1, 22, ICHITA_BLUE_GREY3),     # H1: large, dark
        (2, 15, ICHITA_BLUE_GREY3),     # H2: section headers
        (3, 12, ICHITA_BLUE),           # H3: subsections in brand blue
        (4, 10.5, ICHITA_BLUE),         # H4: minor subsections
    ]
    for level, size, color in heading_configs:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = BRAND_FONT
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.font.bold = True
        # Set Thai font + scaled size on heading style
        h_rPr = hs.element.get_or_add_rPr()
        h_rFonts = h_rPr.find(qn('w:rFonts'))
        if h_rFonts is None:
            h_rFonts = OxmlElement('w:rFonts')
            h_rPr.insert(0, h_rFonts)
        h_rFonts.set(qn('w:cs'), THAI_FONT)
        h_rFonts.set(qn('w:eastAsia'), THAI_FONT)
        h_scaled_hp = str(round(size * 2 * THAI_SCALE))
        h_szCs = h_rPr.find(qn('w:szCs'))
        if h_szCs is not None:
            h_szCs.set(qn('w:val'), h_scaled_hp)
        else:
            h_szCs = OxmlElement('w:szCs')
            h_szCs.set(qn('w:val'), h_scaled_hp)
            h_rPr.append(h_szCs)
        if level == 1:
            hs.paragraph_format.space_before = Pt(20)
            hs.paragraph_format.space_after = Pt(8)
        elif level == 2:
            hs.paragraph_format.space_before = Pt(16)
            hs.paragraph_format.space_after = Pt(6)
        elif level == 3:
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(5)
        else:
            hs.paragraph_format.space_before = Pt(8)
            hs.paragraph_format.space_after = Pt(4)

    # ── List Bullet style ──
    if 'List Bullet' in doc.styles:
        lb = doc.styles['List Bullet']
        lb.font.name = BRAND_FONT
        lb.font.size = Pt(10)
        lb.font.color.rgb = ICHITA_BLUE_GREY3

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Add header and footer ──
    add_header_footer(doc, LOGO_PATH)

    i = 0
    first_h1 = True

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip('\n')

        # ── Horizontal rule (---) ──
        if re.match(r'^---+\s*$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="2978FF"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ── Code block (```) ──
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip('\n').startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            if i < len(lines):
                i += 1
            add_code_block(doc, code_lines)
            continue

        # ── Table (| ... |) ──
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].rstrip('\n').strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1

            if len(table_lines) < 2:
                for tl in table_lines:
                    p = doc.add_paragraph()
                    add_formatted_text(p, tl.strip())
                continue

            header_cells = parse_table_line(table_lines[0])
            data_start = 1
            if len(table_lines) > 1 and is_separator_line(table_lines[1]):
                data_start = 2

            data_rows = []
            for tl in table_lines[data_start:]:
                if not is_separator_line(tl):
                    data_rows.append(parse_table_line(tl))

            add_table_from_rows(doc, header_cells, data_rows)
            continue

        # ── Headings (#, ##, ###, ####) ──
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()

            if level == 1 and first_h1:
                # ── Title: centered, large, Ichita brand ──
                p = doc.add_heading('', level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
                _add_split_run(p, clean_text, BRAND_FONT, Pt(26),
                               ICHITA_BLUE_GREY3, bold=True)
                first_h1 = False
                i += 1

                # Add decorative blue band below title
                add_title_page_band(doc)
                continue
            else:
                p = doc.add_heading('', level=level)
                sizes = {1: Pt(22), 2: Pt(15), 3: Pt(12), 4: Pt(10)}
                colors = {
                    1: ICHITA_BLUE_GREY3,
                    2: ICHITA_BLUE_GREY3,
                    3: ICHITA_BLUE,
                    4: ICHITA_BLUE,
                }

                # For H2, add a left accent bar
                if level == 2:
                    pPr = p._p.get_or_add_pPr()
                    pBdr = parse_xml(
                        f'<w:pBdr {nsdecls("w")}>'
                        f'  <w:left w:val="single" w:sz="24" w:space="6" w:color="2978FF"/>'
                        f'</w:pBdr>'
                    )
                    pPr.append(pBdr)

                add_formatted_text(p, heading_text, base_font=BRAND_FONT,
                                   base_size=sizes[level], base_color=colors[level])
                for run in p.runs:
                    run.font.color.rgb = colors[level]
                    run.font.bold = True

            i += 1
            continue

        # ── Blockquote (> text) ──
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)

            pPr = p._p.get_or_add_pPr()
            # Left border in Ichita Blue
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="2978FF"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

            # Light background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_QUOTE_BG}" w:val="clear"/>')
            pPr.append(shading)

            add_formatted_text(p, quote_text, is_blockquote=True,
                               base_color=ICHITA_BLUE_GREY2)
            i += 1
            continue

        # ── Numbered list (1. item) ──
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if numbered_match:
            num = numbered_match.group(1)
            item_text = numbered_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Number in Ichita Blue
            run = p.add_run(f"{num}. ")
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = ICHITA_BLUE

            add_formatted_text(p, item_text)
            i += 1
            continue

        # ── Bullet point (- item or * item, including nested) ──
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.+)$', stripped)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            bullet_text = bullet_match.group(3)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            if indent_spaces >= 2:
                p.paragraph_format.left_indent = Inches(0.6 + (indent_spaces // 2) * 0.25)

            p.clear()
            add_formatted_text(p, bullet_text)
            i += 1
            continue

        # ── Empty line ──
        if not stripped.strip():
            i += 1
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        add_formatted_text(p, stripped)
        i += 1

    doc.save(output_path)

    # ── Report ──
    size = os.path.getsize(output_path)
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    print(f"Saved: {output_path}")
    print(f"  Size: {size:,} bytes ({size/1024:.1f} KB)")
    print(f"  Paragraphs: {para_count}, Tables: {table_count}")
    print(f"  Font: {BRAND_FONT} + {THAI_FONT} (Thai, {THAI_SCALE}x)")
    print(f"  Brand: ICHITA -- Separation Technologies")


if __name__ == "__main__":
    if len(sys.argv) >= 3:
        input_md = sys.argv[1]
        output_docx = sys.argv[2]
    elif len(sys.argv) == 2:
        input_md = sys.argv[1]
        output_docx = os.path.splitext(input_md)[0] + ".docx"
    else:
        print("Usage: python3 md_to_docx.py input.md [output.docx]")
        sys.exit(1)

    print(f"Converting: {os.path.basename(input_md)}")
    print(f"Brand style: ICHITA")
    convert_md_to_docx(input_md, output_docx)
