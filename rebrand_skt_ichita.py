#!/usr/bin/env python3
"""
Rebrand any DOCX document to Ichita brand identity.

Reads the source DOCX and creates an Ichita-branded version:
- ICHITA logo in header with blue accent line
- Brand fonts (Aeonik → Avenir Next → Calibri)
- Brand colors (#2978FF primary, #263338 headings)
- Dark table headers (#263338), alternating row shading (#EFF2F3)
- Preserves all content, tables (merged cells), images, sections

Usage:
    python3 rebrand_skt_ichita.py input.docx                # → input_ichita.docx
    python3 rebrand_skt_ichita.py input.docx output.docx    # custom output path
    python3 rebrand_skt_ichita.py input.docx --logo my.png  # custom logo
"""

import argparse
import os
import re
import copy
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Ichita Brand Colours ─────────────────────────────────────────────────────

ICHITA_BLUE       = RGBColor(0x29, 0x78, 0xFF)
ICHITA_BLUE_GREY3 = RGBColor(0x26, 0x33, 0x38)

HEX_BLUE      = "2978FF"
HEX_DARK      = "263338"
HEX_TABLE_HDR = "263338"
HEX_TABLE_ALT = "EFF2F3"
HEX_BORDER    = "A0B0B8"
HEX_GREY2     = "788F9C"


# ── Script directory (for resolving relative paths) ─────────────────────────

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


# ── Font Detection ───────────────────────────────────────────────────────────

BRAND_FONT = "Avenir Next"
THAI_FONT  = "Bai Jamjuree"  # Thai font with matched metrics to geometric sans-serif
THAI_SCALE = 0.9              # Thai 9pt / English 10pt — one size down for visual balance

# Check system font dirs + bundled Aeonik-Essentials-Web for Aeonik
_font_dirs = [
    os.path.expanduser("~/Library/Fonts"),
    "/Library/Fonts",
    "/System/Library/Fonts",
    os.path.expanduser("~/.local/share/fonts"),
    os.path.join(SCRIPT_DIR, "Aeonik-Essentials-Web"),
]
for _fd in _font_dirs:
    if os.path.isdir(_fd):
        if any("aeonik" in f.lower() for f in os.listdir(_fd)):
            BRAND_FONT = "Aeonik"
            break


# ── Default Logo Path (resolved relative to script) ─────────────────────────

DEFAULT_LOGO = os.path.join(
    SCRIPT_DIR, "ichita brand ID",
    "ICHITA BRAND BOOK AND COMPANY PROFILE", "LogoV2 3",
    "Digital", "PNG", "Wordmark", "Ichita_Logo-05.png")

# Source document style IDs (common Thai template styles — auto-detected)
# These are checked first; if not matched, detect_heading() does text-based detection.
KNOWN_TITLE_STYLES    = {"0TitleTHSaraban", "Title", "0Title"}
KNOWN_SUBTITLE_STYLES = {"0SubTitleTHSaraban", "Subtitle", "0SubTitle"}
KNOWN_TOPIC_STYLES    = {"1Topic", "Heading1", "Heading 1"}


# ── XML Helpers ──────────────────────────────────────────────────────────────

def get_style_id(p_elem):
    """Get the w:pStyle val from a paragraph element."""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        ps = pPr.find(qn('w:pStyle'))
        if ps is not None:
            return ps.get(qn('w:val'), '')
    return ''


def get_text(p_elem):
    """Get visible text from direct-child w:r/w:t only (avoids MC duplication)."""
    texts = []
    for r in p_elem.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            if t.text:
                texts.append(t.text)
    return ''.join(texts)


def has_images(elem):
    """Check if element contains embedded images (a:blip)."""
    return bool(elem.findall('.//' + qn('a:blip')))


def _is_thai(c):
    """Check if a character is Thai (U+0E00-U+0E7F)."""
    return '\u0e00' <= c <= '\u0e7f'


def _split_thai_latin(text):
    """Split text into segments of (text, is_thai) tuples.
    Groups consecutive Thai chars together, and consecutive non-Thai chars together."""
    if not text:
        return []
    segments = []
    current = text[0]
    current_thai = _is_thai(text[0])
    for c in text[1:]:
        c_thai = _is_thai(c)
        if c_thai == current_thai:
            current += c
        else:
            segments.append((current, current_thai))
            current = c
            current_thai = c_thai
    segments.append((current, current_thai))
    return segments


def _text_is_thai(text):
    """Check if text contains any Thai characters."""
    return any(_is_thai(c) for c in text)


def _text_is_mixed(text):
    """Check if text contains both Thai and non-Thai characters (ignoring whitespace/punctuation)."""
    has_thai = False
    has_latin = False
    for c in text:
        if _is_thai(c):
            has_thai = True
        elif c.isalpha():
            has_latin = True
        if has_thai and has_latin:
            return True
    return False


def split_run_thai_latin(run_elem, parent_elem):
    """Split a mixed Thai/Latin w:r element into separate runs with appropriate fonts.

    For runs containing both Thai and Latin text, splits into multiple <w:r> elements:
    - Thai segments: Bai Jamjuree font at size * THAI_SCALE
    - Latin segments: BRAND_FONT at original size

    Args:
        run_elem: The w:r XML element to potentially split
        parent_elem: The parent element (w:p) containing the run
    """
    # Get the text element
    t_elem = run_elem.find(qn('w:t'))
    if t_elem is None or not t_elem.text:
        return

    text = t_elem.text
    segments = _split_thai_latin(text)

    # If only one segment (all Thai or all Latin), just set the right font
    if len(segments) <= 1:
        return

    # Multiple segments — split the run
    rPr_orig = run_elem.find(qn('w:rPr'))

    # Get current size from rPr
    current_sz = None
    if rPr_orig is not None:
        sz_el = rPr_orig.find(qn('w:sz'))
        if sz_el is not None:
            current_sz = int(sz_el.get(qn('w:val'), '0'))

    # Insert new runs after the original, then remove the original
    insert_after = run_elem
    for seg_text, is_thai in segments:
        new_run = parse_xml(f'<w:r {nsdecls("w")}/>')

        # Copy rPr
        if rPr_orig is not None:
            new_rPr = copy.deepcopy(rPr_orig)
        else:
            new_rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        new_run.insert(0, new_rPr)

        # Set font name
        rf = new_rPr.find(qn('w:rFonts'))
        if rf is None:
            rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            new_rPr.insert(0, rf)
        font = THAI_FONT if is_thai else BRAND_FONT
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rf.set(qn(attr), font)

        # Set size — Thai gets scaled down
        if is_thai and current_sz:
            thai_hp = str(int(current_sz * THAI_SCALE))
            for tag in ('w:sz', 'w:szCs'):
                el = new_rPr.find(qn(tag))
                if el is not None:
                    el.set(qn('w:val'), thai_hp)
                else:
                    new_rPr.append(parse_xml(
                        f'<{tag} {nsdecls("w")} w:val="{thai_hp}"/>'))

        # Add text element
        new_t = parse_xml(f'<w:t {nsdecls("w")}/>')
        new_t.text = seg_text
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_run.append(new_t)

        # Insert after current position
        insert_after.addnext(new_run)
        insert_after = new_run

    # Remove the original run
    parent_elem.remove(run_elem)


def set_font(run_elem, font_name=None, size_pt=None, color_hex=None, bold=None):
    """Set font properties on a w:r element at the XML level.

    Detects Thai text and sets Bai Jamjuree font with scaled size.
    For Latin text, uses BRAND_FONT at the specified size.
    Mixed runs are handled later by split_run_thai_latin().
    """
    if font_name is None:
        font_name = BRAND_FONT

    rPr = run_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        run_elem.insert(0, rPr)

    # ── Detect if this run's text is Thai ──
    t_elem = run_elem.find(qn('w:t'))
    run_text = t_elem.text if t_elem is not None and t_elem.text else ""
    is_thai = _text_is_thai(run_text) and not _text_is_mixed(run_text)

    # ── Font name — Thai gets Bai Jamjuree, Latin gets BRAND_FONT ──
    actual_font = THAI_FONT if is_thai else font_name
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(attr), actual_font)

    # ── Size (half-points) ──
    # Thai text gets scaled down by THAI_SCALE for visual balance
    if size_pt is not None:
        if is_thai:
            hp = str(int(size_pt * THAI_SCALE * 2))
        else:
            hp = str(int(size_pt * 2))
        for tag in ('w:sz', 'w:szCs'):
            el = rPr.find(qn(tag))
            if el is not None:
                el.set(qn('w:val'), hp)
            else:
                new_el = parse_xml(f'<{tag} {nsdecls("w")} w:val="{hp}"/>')
                rPr.append(new_el)

    # ── Colour ──
    if color_hex is not None:
        el = rPr.find(qn('w:color'))
        if el is not None:
            el.set(qn('w:val'), color_hex)
        else:
            el = parse_xml(
                f'<w:color {nsdecls("w")} w:val="{color_hex}"/>')
            rPr.append(el)

    # ── Bold ──
    if bold is not None:
        for tag in ('w:b', 'w:bCs'):
            el = rPr.find(qn(tag))
            if bold:
                if el is None:
                    rPr.append(parse_xml(f'<{tag} {nsdecls("w")}/>'))
            else:
                if el is not None:
                    rPr.remove(el)


def set_cell_shading(tc_elem, color_hex):
    """Apply background shading to a table cell XML element."""
    tcPr = tc_elem.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        tc_elem.insert(0, tcPr)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shading)


def ensure_pPr(p_elem):
    """Ensure paragraph has a pPr element; return it."""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        p_elem.insert(0, pPr)
    return pPr


def add_left_accent(p_elem, color="2978FF", sz="24", space="6"):
    """Add a left blue accent bar to a paragraph."""
    pPr = ensure_pPr(p_elem)
    old = pPr.find(qn('w:pBdr'))
    if old is not None:
        pPr.remove(old)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="{space}"'
        f'         w:color="{color}"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)


def add_bottom_band(p_elem, color="2978FF", sz="24"):
    """Add a thick bottom blue band below a paragraph (title decoration)."""
    pPr = ensure_pPr(p_elem)
    old = pPr.find(qn('w:pBdr'))
    if old is not None:
        pPr.remove(old)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="1"'
        f'            w:color="{color}"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)


def set_alignment(p_elem, val='center'):
    """Set paragraph alignment (center, left, right, both)."""
    pPr = ensure_pPr(p_elem)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = parse_xml(f'<w:jc {nsdecls("w")}/>')
        pPr.append(jc)
    jc.set(qn('w:val'), val)


# ── Heading Detection ────────────────────────────────────────────────────────

def detect_heading(text):
    """Detect heading type from Normal-styled paragraph text."""
    text = text.strip()
    if not text:
        return None

    # Figure/table captions
    if text.startswith('รูปที่') or text.startswith('ตารางที่'):
        return 'caption'

    # "Appendix" heading
    if text.lower() == 'appendix':
        return 'section'

    # Appendix sub-headings: A. ..., B. ..., C. ..., D. ...
    if re.match(r'^[A-D]\.?\s*\S', text) and len(text) < 80:
        return 'subsection'

    # Section headings: 1. ..., 2. ..., etc. (short text)
    if re.match(r'^\d+\.\s+\S', text) and len(text) < 80:
        return 'section'

    # Subsection headings: 3.1 ..., 3.2 ..., etc. (short text)
    if re.match(r'^\d+\.\d+\s+\S', text) and len(text) < 80:
        return 'subsection'

    return None


# ── Paragraph Styling ────────────────────────────────────────────────────────

def style_runs(p_elem, size_pt=10.5, color_hex="263338", bold=None):
    """Apply brand font to all text runs in a paragraph (skip image runs).
    After setting fonts, splits mixed Thai/Latin runs into separate elements."""
    for run in p_elem.findall('.//' + qn('w:r')):
        if has_images(run):
            continue
        set_font(run, size_pt=size_pt, color_hex=color_hex, bold=bold)

    # Second pass: split any mixed Thai/Latin runs into separate elements
    for run in list(p_elem.findall('.//' + qn('w:r'))):
        if has_images(run):
            continue
        t_elem = run.find(qn('w:t'))
        if t_elem is not None and t_elem.text and _text_is_mixed(t_elem.text):
            split_run_thai_latin(run, p_elem)


def style_paragraph(p_elem, style_id, text):
    """Apply Ichita brand styling to a paragraph."""

    # ── Title ──
    if style_id in KNOWN_TITLE_STYLES:
        style_runs(p_elem, size_pt=26, color_hex=HEX_DARK, bold=True)
        set_alignment(p_elem, 'center')
        add_bottom_band(p_elem)
        return

    # ── Subtitle ──
    if style_id in KNOWN_SUBTITLE_STYLES:
        style_runs(p_elem, size_pt=15, color_hex=HEX_GREY2)
        set_alignment(p_elem, 'center')
        return

    # ── Topic (H2 from source template) ──
    if style_id in KNOWN_TOPIC_STYLES:
        style_runs(p_elem, size_pt=15, color_hex=HEX_DARK, bold=True)
        add_left_accent(p_elem)
        pPr = ensure_pPr(p_elem)
        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:before'), '280')
        sp.set(qn('w:after'), '120')
        return

    # ── Detect heading by text content (Normal-styled headings) ──
    heading = detect_heading(text)

    if heading == 'section':
        style_runs(p_elem, size_pt=14, color_hex=HEX_DARK, bold=True)
        add_left_accent(p_elem)
        pPr = ensure_pPr(p_elem)
        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
        # Standardise spacing (override source inconsistencies)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:before'), '280')
        sp.set(qn('w:after'), '120')
    elif heading == 'subsection':
        style_runs(p_elem, size_pt=12, color_hex=HEX_BLUE, bold=True)
        pPr = ensure_pPr(p_elem)
        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:before'), '200')
        sp.set(qn('w:after'), '120')
    elif heading == 'caption':
        style_runs(p_elem, size_pt=10.5, color_hex=HEX_GREY2)
        set_alignment(p_elem, 'center')
        pPr = ensure_pPr(p_elem)
        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:before'), '120')
        sp.set(qn('w:after'), '60')

    else:
        # Normal body text (12pt — one size up for Thai readability)
        style_runs(p_elem, size_pt=12, color_hex=HEX_DARK)
        # Ensure consistent paragraph spacing — reset both before and after
        # to override any values inherited from the source document.
        pPr = ensure_pPr(p_elem)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '120')


# ── Table Styling ────────────────────────────────────────────────────────────

def style_table(tbl_elem):
    """Apply Ichita brand styling to a table XML element.

    Handles: dark header rows, alternating data row shading, brand borders,
    multi-row headers (vMerge), and image-only tables (no header styling).
    """
    # ── Table-level borders ──
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl_elem.insert(0, tblPr)

    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0"'
        f'         w:color="{HEX_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0"'
        f'          w:color="{HEX_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0"'
        f'             w:color="{HEX_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0"'
        f'           w:color="{HEX_BORDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0"'
        f'              w:color="{HEX_BORDER}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0"'
        f'              w:color="{HEX_BORDER}"/>'
        f'</w:tblBorders>')
    tblPr.append(borders)

    rows = tbl_elem.findall(qn('w:tr'))
    if not rows:
        return

    # ── Detect image-only table (photo grid — skip header styling) ──
    r0_cells = rows[0].findall(qn('w:tc'))
    is_image_table = (
        r0_cells and all(has_images(tc) for tc in r0_cells)
    )

    # ── Detect multi-row header (vMerge=restart in row 0) ──
    header_rows = 1
    if not is_image_table:
        for tc in r0_cells:
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                vm = tcPr.find(qn('w:vMerge'))
                if vm is not None and vm.get(qn('w:val')) == 'restart':
                    header_rows = 2
                    break

    # ── Detect wide tables (>10 cols) — use smaller font ──
    num_cols = len(rows[0].findall(qn('w:tc'))) if rows else 0
    font_sz = 8 if num_cols > 10 else 10
    hdr_sz = 8 if num_cols > 10 else 10

    # ── Style each row ──
    for ri, tr in enumerate(rows):
        # For wide tables, set row height to "atLeast" so text isn't clipped
        if num_cols > 10:
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
                tr.insert(0, trPr)
            th = trPr.find(qn('w:trHeight'))
            if th is not None:
                th.set(qn('w:hRule'), 'atLeast')
                # Ensure minimum height for Thai text
                cur = int(th.get(qn('w:val'), '0'))
                if cur < 320:
                    th.set(qn('w:val'), '320')
            else:
                trPr.append(parse_xml(
                    f'<w:trHeight {nsdecls("w")}'
                    f' w:val="320" w:hRule="atLeast"/>'))

        for tc in tr.findall(qn('w:tc')):
            if is_image_table:
                # Image table: brand font only, no header/alt styling
                for run in tc.findall('.//' + qn('w:r')):
                    if not has_images(run):
                        set_font(run, size_pt=font_sz, color_hex=HEX_DARK)

            elif ri < header_rows:
                # Header row(s): dark bg, white bold text
                set_cell_shading(tc, HEX_TABLE_HDR)
                for run in tc.findall('.//' + qn('w:r')):
                    if not has_images(run):
                        set_font(run, size_pt=hdr_sz, color_hex="FFFFFF",
                                 bold=True)

            else:
                # Data rows: alternating shading, clear source colors
                data_idx = ri - header_rows
                if data_idx % 2 == 1:
                    set_cell_shading(tc, HEX_TABLE_ALT)
                else:
                    set_cell_shading(tc, "FFFFFF")
                for run in tc.findall('.//' + qn('w:r')):
                    if not has_images(run):
                        set_font(run, size_pt=font_sz, color_hex=HEX_DARK)


def squeeze_wide_tables(body, usable_twips):
    """Scale tables that exceed usable_twips to fit within the page width.

    Proportionally reduces gridCol widths and tcW cell widths so the table
    fits without losing column proportions or content integrity.
    """
    for tbl in body.iter(qn('w:tbl')):
        grid = tbl.find(qn('w:tblGrid'))
        if grid is None:
            continue
        grid_cols = grid.findall(qn('w:gridCol'))
        gc_widths = []
        for gc in grid_cols:
            w = gc.get(qn('w:w'))
            gc_widths.append(int(w) if w else 0)
        total = sum(gc_widths)
        if total <= usable_twips or total == 0:
            continue

        # Scale factor
        factor = usable_twips / total
        print(f"  Squeeze table: {total} → {usable_twips} twips "
              f"(factor {factor:.3f}, {len(grid_cols)} grid cols)")

        # Scale grid columns
        for gc, old_w in zip(grid_cols, gc_widths):
            gc.set(qn('w:w'), str(int(old_w * factor)))

        # Scale tblW if absolute
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                wtype = tblW.get(qn('w:type'), 'auto')
                if wtype == 'dxa':
                    old_tw = int(tblW.get(qn('w:w'), '0'))
                    tblW.set(qn('w:w'), str(int(old_tw * factor)))

        # Scale cell widths in every row
        for tr in tbl.findall(qn('w:tr')):
            for tc in tr.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    continue
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    wtype = tcW.get(qn('w:type'), 'dxa')
                    if wtype == 'dxa':
                        old_cw = int(tcW.get(qn('w:w'), '0'))
                        tcW.set(qn('w:w'), str(int(old_cw * factor)))


# ── Image Relationship Copy ──────────────────────────────────────────────────

def copy_image_rels(src_part, dst_part, elem):
    """Copy image relationships from source to destination for all blips."""
    for blip in elem.findall('.//' + qn('a:blip')):
        rId = blip.get(qn('r:embed'))
        if rId and rId in src_part.rels:
            rel = src_part.rels[rId]
            new_rId = dst_part.relate_to(rel.target_part, rel.reltype)
            blip.set(qn('r:embed'), new_rId)


# ── Header / Footer ─────────────────────────────────────────────────────────

def add_header_footer(doc, logo_path):
    """Add ICHITA logo header and clean footer to all sections."""
    for section in doc.sections:
        # ── Header ──
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()

        hp = (header.paragraphs[0]
              if header.paragraphs else header.add_paragraph())
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(4)

        if os.path.exists(logo_path):
            run = hp.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
        else:
            run = hp.add_run("ICHITA\u2122")
            run.font.name = BRAND_FONT
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = ICHITA_BLUE_GREY3

        # Blue accent line under header
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="4"'
            f'            w:color="{HEX_BLUE}"/>'
            f'</w:pBdr>')
        pPr.append(pBdr)

        # ── Footer (clear) ──
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()


# ── Cleanup ──────────────────────────────────────────────────────────────────

def cleanup_empty_space(body):
    """Remove title-page logo, 'ICHITA's solutions', and excess empty paras.

    Walks direct children of body. Removes:
    - The big ICHITA circle logo paragraph (image-only, before the title)
    - The "ICHITA's solutions" tagline paragraph
    - Consecutive empty paragraphs beyond the first in any run
    """
    found_title = False
    removed = 0

    # Collect direct-child paragraphs (not inside tables/sectPr)
    children = list(body)
    prev_was_empty = False

    for elem in children:
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            prev_was_empty = False
            continue

        text = get_text(elem).strip()
        has_img = has_images(elem)
        style_id = get_style_id(elem)

        # Track when we've passed the title (by style or by content heading)
        if style_id in KNOWN_TITLE_STYLES | KNOWN_TOPIC_STYLES or 'รายงานผลการทดสอบ' in text:
            found_title = True

        # Remove big logo image paragraph (image-only, before title)
        if not found_title and has_img and not text:
            body.remove(elem)
            removed += 1
            continue

        # Remove "ICHITA's solutions" tagline
        if "ICHITA" in text and "solution" in text.lower():
            body.remove(elem)
            removed += 1
            continue

        # Collapse consecutive empty paragraphs (keep max 1)
        is_empty = (not text and not has_img)
        # Don't remove paragraphs that carry a section break
        has_sect = False
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            has_sect = True

        if is_empty and not has_sect:
            if prev_was_empty:
                body.remove(elem)
                removed += 1
                continue
            prev_was_empty = True
        else:
            prev_was_empty = False

    print(f"  Cleanup: removed {removed} empty/logo paragraphs")


# ── Title Page Redesign ──────────────────────────────────────────────────────

def make_para(text="", size_pt=12, color_hex=HEX_DARK, bold=False,
              align='left', space_before=0, space_after=0):
    """Create a new w:p element with formatted text and spacing.
    Splits mixed Thai/Latin text into separate runs."""
    p = parse_xml(f'<w:p {nsdecls("w")}/>')

    if text:
        # Split into Thai/Latin segments for proper font assignment
        segments = _split_thai_latin(text)
        for seg_text, is_thai in segments:
            r = parse_xml(f'<w:r {nsdecls("w")}/>')
            t = parse_xml(f'<w:t {nsdecls("w")}/>')
            t.text = seg_text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            font = THAI_FONT if is_thai else BRAND_FONT
            sz = size_pt * THAI_SCALE if is_thai else size_pt
            set_font(r, font_name=font, size_pt=sz, color_hex=color_hex, bold=bold)
            p.append(r)

    pPr = ensure_pPr(p)
    if align != 'left':
        pPr.append(parse_xml(
            f'<w:jc {nsdecls("w")} w:val="{align}"/>'))
    sb = str(int(space_before * 20))   # pt → twips
    sa = str(int(space_after * 20))
    pPr.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="{sb}" w:after="{sa}"/>'))
    return p


def create_meta_table(items):
    """Build a clean 2-column metadata table (label | value).

    Horizontal dividers only, no vertical lines — modern, professional look.
    Uses Bai Jamjuree for Thai labels, BRAND_FONT for Latin values.
    """
    label_sz = str(int(11 * THAI_SCALE * 2))  # Thai scaled size in half-points
    value_sz = "22"                             # 11pt in half-points

    rows_xml = ""
    for label, value in items:
        # Label: Thai text → Bai Jamjuree at scaled size
        label_font = THAI_FONT if _text_is_thai(label) else BRAND_FONT
        label_hp = label_sz if _text_is_thai(label) else value_sz

        # Value: detect Thai content
        value_font = THAI_FONT if _text_is_thai(value) else BRAND_FONT
        value_hp = label_sz if _text_is_thai(value) else value_sz

        rows_xml += (
            '<w:tr>'
            '  <w:tc>'
            '    <w:tcPr><w:tcW w:w="2600" w:type="dxa"/></w:tcPr>'
            '    <w:p><w:pPr>'
            '      <w:spacing w:before="50" w:after="50"/>'
            '    </w:pPr>'
            '    <w:r><w:rPr>'
            f'      <w:rFonts w:ascii="{label_font}" w:hAnsi="{label_font}"'
            f'               w:cs="{label_font}" w:eastAsia="{label_font}"/>'
            f'      <w:sz w:val="{label_hp}"/><w:szCs w:val="{label_hp}"/>'
            f'      <w:color w:val="{HEX_BLUE}"/>'
            '      <w:b/><w:bCs/>'
            f'    </w:rPr><w:t xml:space="preserve">{label}</w:t></w:r>'
            '    </w:p>'
            '  </w:tc>'
            '  <w:tc>'
            '    <w:tcPr><w:tcW w:w="7400" w:type="dxa"/></w:tcPr>'
            '    <w:p><w:pPr>'
            '      <w:spacing w:before="50" w:after="50"/>'
            '    </w:pPr>'
            '    <w:r><w:rPr>'
            f'      <w:rFonts w:ascii="{value_font}" w:hAnsi="{value_font}"'
            f'               w:cs="{value_font}" w:eastAsia="{value_font}"/>'
            f'      <w:sz w:val="{value_hp}"/><w:szCs w:val="{value_hp}"/>'
            f'      <w:color w:val="{HEX_DARK}"/>'
            f'    </w:rPr><w:t xml:space="preserve">{value}</w:t></w:r>'
            '    </w:p>'
            '  </w:tc>'
            '</w:tr>'
        )

    tbl_xml = (
        f'<w:tbl {nsdecls("w")}>'
        '  <w:tblPr>'
        '    <w:tblW w:w="5000" w:type="pct"/>'
        '    <w:jc w:val="center"/>'
        '    <w:tblBorders>'
        f'      <w:top w:val="single" w:sz="6" w:space="0"'
        f'            w:color="{HEX_BLUE}"/>'
        f'      <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'      <w:bottom w:val="single" w:sz="6" w:space="0"'
        f'               w:color="{HEX_BLUE}"/>'
        f'      <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'      <w:insideH w:val="single" w:sz="2" w:space="0"'
        f'                 w:color="{HEX_BORDER}"/>'
        f'      <w:insideV w:val="none" w:sz="0" w:space="0"'
        f'                 w:color="auto"/>'
        '    </w:tblBorders>'
        '    <w:tblCellMar>'
        '      <w:top w:w="60" w:type="dxa"/>'
        '      <w:left w:w="120" w:type="dxa"/>'
        '      <w:bottom w:w="60" w:type="dxa"/>'
        '      <w:right w:w="120" w:type="dxa"/>'
        '    </w:tblCellMar>'
        '  </w:tblPr>'
        '  <w:tblGrid>'
        '    <w:gridCol w:w="2600"/>'
        '    <w:gridCol w:w="7400"/>'
        '  </w:tblGrid>'
        f'  {rows_xml}'
        '</w:tbl>'
    )
    return parse_xml(tbl_xml)


def redesign_title_page(body):
    """Replace old title page with a professional Ichita-branded cover.

    Extracts metadata (strips English translations), title, and subtitle.
    Builds: Title → blue band → subtitle → metadata table → section break.
    """
    metadata = {}
    title_text = ""
    subtitle_text = ""
    sect_break_xml = None
    to_remove = []

    # ── Extract data from existing elements ──
    for elem in list(body):
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue

        text = get_text(elem).strip()
        style_id = get_style_id(elem)

        # Metadata fields — extract value, strip (English Translation)
        for thai_key, dict_key in [
            ('ชื่อโครงการ', 'project'),
            ('ชื่องาน', 'case'),
            ('ลูกค้า', 'customer'),
            ('วันที่จัดทำ', 'date'),
            ('หมายเหตุ', 'notes'),
        ]:
            if thai_key in text:
                clean = re.sub(r'\s*\([^)]*\)\s*', ' ', text)
                _, _, value = clean.partition(':')
                metadata[dict_key] = value.strip()

                # Preserve section break if this paragraph has one
                pPr = elem.find(qn('w:pPr'))
                if pPr is not None:
                    sect = pPr.find(qn('w:sectPr'))
                    if sect is not None:
                        sect_break_xml = copy.deepcopy(sect)

                to_remove.append(elem)
                break

        # Title and subtitle — detect by style OR by text content
        if style_id in KNOWN_TITLE_STYLES or (not title_text and 'รายงานผลการทดสอบ' in text):
            title_text = text
            to_remove.append(elem)
        elif style_id in KNOWN_SUBTITLE_STYLES or (not subtitle_text and 'ระบบกำจัดสี' in text):
            subtitle_text = text
            to_remove.append(elem)

    # Also remove remaining empty paragraphs before first Topic heading
    for elem in list(body):
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        text = get_text(elem).strip()
        style_id = get_style_id(elem)
        if elem in to_remove:
            continue
        # Stop at first Topic heading (real content starts)
        if style_id in KNOWN_TOPIC_STYLES:
            break
        # Skip metadata / title / subtitle (already handled)
        if any(k in text for k in ['ชื่อ', 'ลูกค้า', 'วันที่', 'หมายเหตุ',
                                     'รายงานผล', 'ระบบกำจัดสี']):
            continue
        if not text and not has_images(elem):
            pPr = elem.find(qn('w:pPr'))
            if pPr is None or pPr.find(qn('w:sectPr')) is None:
                to_remove.append(elem)

    # ── Remove old elements ──
    for elem in to_remove:
        if elem.getparent() is body:
            body.remove(elem)

    # ── Find insertion point (first remaining content element) ──
    insert_before = None
    for child in body:
        if child.tag != qn('w:sectPr'):
            insert_before = child
            break

    def insert(elem):
        if insert_before is not None:
            insert_before.addprevious(elem)
        else:
            final = body.find(qn('w:sectPr'))
            if final is not None:
                final.addprevious(elem)
            else:
                body.append(elem)

    # ── Build new title page ──

    # 1. Top spacer
    insert(make_para(space_after=80))

    # 2. Title (very large, centered, bold)
    insert(make_para(title_text, size_pt=36, color_hex=HEX_DARK,
                     bold=True, align='center', space_after=10))

    # 3. Blue accent band
    band = make_para(space_before=0, space_after=14)
    add_bottom_band(band, color=HEX_BLUE, sz="24")
    insert(band)

    # 4. Subtitle (centered, grey)
    insert(make_para(subtitle_text, size_pt=16, color_hex=HEX_GREY2,
                     align='center', space_before=8, space_after=36))

    # 5. Metadata table
    meta_items = [
        ('ชื่อโครงการ',       metadata.get('project', '')),
        ('ชื่องาน',           metadata.get('case', '')),
        ('ลูกค้า',            metadata.get('customer', '')),
        ('วันที่จัดทำ',        metadata.get('date', '')),
        ('หมายเหตุโครงการ',   metadata.get('notes', '')),
    ]
    insert(create_meta_table(meta_items))

    # 6. Section break paragraph (page break between cover and content)
    sect_para = make_para(space_before=0, space_after=0)
    if sect_break_xml is not None:
        pPr = ensure_pPr(sect_para)
        pPr.append(sect_break_xml)
    insert(sect_para)

    print(f"  Title page: redesigned ({len(meta_items)} metadata fields)")


# ── Main Rebranding ──────────────────────────────────────────────────────────

def rebrand(source, output, logo=None):
    """Open source DOCX, create Ichita-branded copy.

    Args:
        source: Path to source DOCX file
        output: Path to write branded output DOCX
        logo: Path to logo PNG (default: bundled Ichita wordmark)
    """
    if logo is None:
        logo = DEFAULT_LOGO

    print(f"Source: {os.path.basename(source)}")
    print(f"Output: {os.path.basename(output)}")
    print(f"Font:   {BRAND_FONT}")

    src_doc = Document(source)
    dst_doc = Document()

    src_body = src_doc.element.body
    dst_body = dst_doc.element.body

    # Clear destination body (keep final sectPr)
    for child in list(dst_body):
        if child.tag != qn('w:sectPr'):
            dst_body.remove(child)

    # ── Deep-copy all source body children ──
    src_part = src_doc.part
    dst_part = dst_doc.part

    for child in src_body:
        tag = child.tag.split('}')[-1]
        new_elem = copy.deepcopy(child)

        # Copy image relationships
        copy_image_rels(src_part, dst_part, new_elem)

        if tag == 'sectPr':
            # Replace destination's final sectPr
            old = dst_body.find(qn('w:sectPr'))
            if old is not None:
                dst_body.remove(old)
            dst_body.append(new_elem)
        else:
            # Insert before final sectPr
            final = dst_body.find(qn('w:sectPr'))
            if final is not None:
                final.addprevious(new_elem)
            else:
                dst_body.append(new_elem)

    # ── Remove source header/footer references (we add our own) ──
    for sectPr in dst_body.iter(qn('w:sectPr')):
        for ref in list(sectPr.findall(qn('w:headerReference'))):
            sectPr.remove(ref)
        for ref in list(sectPr.findall(qn('w:footerReference'))):
            sectPr.remove(ref)
        # Remove titlePg — source doc has "different first page" enabled,
        # which hides our header on the first page of each section
        title_pg = sectPr.find(qn('w:titlePg'))
        if title_pg is not None:
            sectPr.remove(title_pg)

    # ── Remove logo, tagline, and excess empty space ──
    cleanup_empty_space(dst_body)

    # (margins set at the end, after all section modifications)

    # ── Restyle paragraphs ──
    para_count = 0
    for p in dst_body.iter(qn('w:p')):
        style_id = get_style_id(p)
        text = get_text(p)
        style_paragraph(p, style_id, text)
        para_count += 1

    # ── Restyle tables ──
    tbl_count = 0
    for tbl in dst_body.iter(qn('w:tbl')):
        style_table(tbl)
        tbl_count += 1

    # ── Reorder image tables: caption above table, text after ──
    # Original order:  text → img table → caption
    # New order:        caption → img table → text
    # Caption on top keeps it glued to its table regardless of pagination,
    # and the text paragraph flows naturally after without being split.
    children = list(dst_body)
    for i, child in enumerate(children):
        if child.tag.split('}')[-1] != 'tbl':
            continue
        if not child.findall('.//' + qn('w:drawing')):
            continue
        # Check next sibling is a caption (starts with "รูปที่")
        caption = None
        if i + 1 < len(children) and children[i + 1].tag == qn('w:p'):
            nxt = children[i + 1]
            cap_text = ''.join(t.text or '' for t in nxt.findall('.//' + qn('w:t'))).strip()
            if cap_text.startswith('รูปที่'):
                caption = nxt
        if caption is not None:
            # Move caption to just before the image table
            child.addprevious(caption)

    # ── Enforce consistent spacing between tables and adjacent text ──
    # Tables in the source often have 0 spacing to adjacent paragraphs,
    # which looks cramped with the smaller branded fonts.
    SPACE_AROUND_TABLE = "160"  # 8pt in twips (half-points * 10)
    children = list(dst_body)
    for i, child in enumerate(children):
        if child.tag.split('}')[-1] != 'tbl':
            continue

        # Paragraph BEFORE table: ensure space_after
        if i > 0 and children[i - 1].tag == qn('w:p'):
            prev_p = children[i - 1]
            pPr = ensure_pPr(prev_p)
            sp = pPr.find(qn('w:spacing'))
            if sp is None:
                sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
                pPr.append(sp)
            cur_after = int(sp.get(qn('w:after'), '0'))
            if cur_after < int(SPACE_AROUND_TABLE):
                sp.set(qn('w:after'), SPACE_AROUND_TABLE)

        # Paragraph AFTER table: ensure space_before
        if i + 1 < len(children) and children[i + 1].tag == qn('w:p'):
            next_p = children[i + 1]
            pPr = ensure_pPr(next_p)
            sp = pPr.find(qn('w:spacing'))
            if sp is None:
                sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
                pPr.append(sp)
            cur_before = int(sp.get(qn('w:before'), '0'))
            if cur_before < int(SPACE_AROUND_TABLE):
                sp.set(qn('w:before'), SPACE_AROUND_TABLE)

    # ── Redesign title page (after restyling so it won't be overridden) ──
    redesign_title_page(dst_body)

    # ── Remove empty paragraphs in appendix area ──
    # These cause blank pages between appendix tables
    in_appendix = False
    for child in list(dst_body):
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            text = get_text(child).strip()
            if text.lower() == 'appendix' or text.startswith('A.'):
                in_appendix = True
            if in_appendix and not text and not has_images(child):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    continue  # keep section-break paragraphs
                dst_body.remove(child)

    # ── Remove empty trailing section ──
    # If the last inline sectPr creates a section with no content,
    # merge it: copy its orientation to the final sectPr and remove it
    children = list(dst_body)
    final_sect = dst_body.find(qn('w:sectPr'))
    if final_sect is not None:
        # Find last inline section break
        last_sb = None
        last_sb_idx = None
        for i, child in enumerate(children):
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    last_sb = child
                    last_sb_idx = i
        # Check if there's no content between last section break and final sectPr
        if last_sb is not None:
            has_content = False
            for sib in children[last_sb_idx + 1:]:
                if sib.tag == qn('w:sectPr'):
                    continue
                if sib.tag == qn('w:tbl'):
                    has_content = True
                    break
                if sib.tag == qn('w:p'):
                    t = get_text(sib).strip()
                    if t or has_images(sib):
                        has_content = True
                        break
            if not has_content:
                # Copy page size/orientation from inline sectPr to final
                pPr = last_sb.find(qn('w:pPr'))
                inline_sect = pPr.find(qn('w:sectPr'))
                src_pgSz = inline_sect.find(qn('w:pgSz'))
                dst_pgSz = final_sect.find(qn('w:pgSz'))
                if src_pgSz is not None and dst_pgSz is not None:
                    for attr in (qn('w:w'), qn('w:h'), qn('w:orient')):
                        val = src_pgSz.get(attr)
                        if val is not None:
                            dst_pgSz.set(attr, val)
                        elif attr == qn('w:orient'):
                            # Remove orient attr if not present (portrait)
                            if attr in dst_pgSz.attrib:
                                del dst_pgSz.attrib[attr]
                dst_body.remove(last_sb)

    # ── Set page margins (LAST — after all section modifications) ──
    LANDSCAPE_MARGIN = Cm(1.0)   # tight side margins for wide appendix tables
    LANDSCAPE_TOP    = Cm(3.0)   # top margin consistent with portrait header gap
    PORTRAIT_MARGIN  = Cm(2.5)
    PORTRAIT_TOP     = Cm(3.5)   # extra top margin: header + gap before body text
    for section in dst_doc.sections:
        if section.orientation:  # Landscape
            section.top_margin    = LANDSCAPE_TOP
            section.bottom_margin = LANDSCAPE_MARGIN
            section.left_margin   = LANDSCAPE_MARGIN
            section.right_margin  = LANDSCAPE_MARGIN
        else:  # Portrait
            section.top_margin    = PORTRAIT_TOP
            section.bottom_margin = PORTRAIT_MARGIN
            section.left_margin   = PORTRAIT_MARGIN
            section.right_margin  = PORTRAIT_MARGIN

    # ── Squeeze wide tables to fit page ──
    # A4 landscape = 16838 twips wide; with 1 cm margins → usable ~15,700
    landscape_usable = int((Cm(29.7) - 2 * LANDSCAPE_MARGIN) / 914400 * 1440)
    squeeze_wide_tables(dst_body, landscape_usable)

    # ── Page-break control for headings before tables ──
    # Remove ALL pageBreakBefore first, then add back only where needed
    for p in dst_body.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                pPr.remove(pb)

    # Add pageBreakBefore to keep headings/captions with their tables
    # NOT appendix A/B — A flows from content; B follows section break
    children = list(dst_body)
    for i, child in enumerate(children):
        if child.tag != qn('w:p'):
            continue
        text = get_text(child).strip()

        target = None  # element to receive pageBreakBefore

        # "Appendix" section heading — start on new page with A.
        if text.lower() == 'appendix':
            target = child

        # Appendix C heading — force new page so heading stays with table
        if re.match(r'^C\.', text):
            target = child

        # Table captions: no forced page break — keepWithNext on caption
        # and heading above naturally keeps them together with the table

        if target is not None:
            pPr = ensure_pPr(target)
            # Don't add duplicate
            if pPr.find(qn('w:pageBreakBefore')) is None:
                pPr.append(parse_xml(
                    f'<w:pageBreakBefore {nsdecls("w")}/>'))
                t = get_text(target).strip()
                print(f"  PB_BEFORE: {t[:60]}")

    # ── Zero space_before on page-top paragraphs ──
    # Paragraphs that start at the top of a page (pageBreakBefore or
    # first after a section break) should have space_before=0 so the
    # gap from header to first content line is consistent everywhere.
    children = list(dst_body)
    for i, child in enumerate(children):
        if child.tag != qn('w:p'):
            continue
        pPr = child.find(qn('w:pPr'))
        if pPr is None:
            continue

        zero_it = False

        # Case 1: paragraph has pageBreakBefore
        if pPr.find(qn('w:pageBreakBefore')) is not None:
            zero_it = True

        # Case 2: first paragraph after a section break
        if i > 0 and children[i - 1].tag == qn('w:p'):
            prev_pPr = children[i - 1].find(qn('w:pPr'))
            if prev_pPr is not None and prev_pPr.find(qn('w:sectPr')) is not None:
                zero_it = True

        if zero_it:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                sp.set(qn('w:before'), '0')

    # ── Document default style ──
    style = dst_doc.styles['Normal']
    style.font.name = BRAND_FONT
    style.font.size = Pt(12)
    style.font.color.rgb = ICHITA_BLUE_GREY3
    # Set Thai (Complex Script) font + scaled size on default style
    n_rPr = style.element.get_or_add_rPr()
    n_rFonts = n_rPr.find(qn('w:rFonts'))
    if n_rFonts is None:
        n_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        n_rPr.insert(0, n_rFonts)
    n_rFonts.set(qn('w:cs'), THAI_FONT)
    n_rFonts.set(qn('w:eastAsia'), THAI_FONT)

    # ── ICHITA header/footer on every page ──
    add_header_footer(dst_doc, logo)

    # ── Save ──
    dst_doc.save(output)

    # ── Report ──
    size = os.path.getsize(output)
    img_count = sum(
        1 for _ in dst_body.iter(qn('a:blip'))
    )
    section_count = len(dst_doc.sections)
    print(f"\n{'='*60}")
    print(f"  Saved: {output}")
    print(f"  Size:       {size:,} bytes ({size/1024/1024:.1f} MB)")
    print(f"  Paragraphs: {para_count}")
    print(f"  Tables:     {tbl_count}")
    print(f"  Images:     {img_count} blip references")
    print(f"  Sections:   {section_count}")
    print(f"  Font:       {BRAND_FONT} + {THAI_FONT} (Thai, {THAI_SCALE}x)")
    print(f"  Brand:      ICHITA -- Separation Technologies")
    print(f"{'='*60}")


def main():
    parser = argparse.ArgumentParser(
        description="Rebrand a DOCX document to Ichita brand identity.",
        epilog="Examples:\n"
               "  python3 rebrand_skt_ichita.py report.docx\n"
               "  python3 rebrand_skt_ichita.py report.docx branded.docx\n"
               "  python3 rebrand_skt_ichita.py report.docx --logo custom_logo.png\n",
        formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("input", help="Source DOCX file to rebrand")
    parser.add_argument("output", nargs="?", default=None,
                        help="Output DOCX path (default: <input>_ichita.docx)")
    parser.add_argument("--logo", default=None,
                        help="Path to logo PNG for header "
                             f"(default: bundled Ichita wordmark)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: source file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    output = args.output
    if output is None:
        base, ext = os.path.splitext(args.input)
        output = f"{base}_ichita{ext}"

    logo = args.logo
    if logo and not os.path.exists(logo):
        print(f"Warning: logo file not found: {logo} — will use text fallback",
              file=sys.stderr)

    rebrand(args.input, output, logo=logo)


if __name__ == "__main__":
    main()
